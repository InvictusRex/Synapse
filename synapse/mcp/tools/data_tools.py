"""
Data Processing Tools - Handle CSV, Excel, Word, JSON files
"""
import os
import json
from typing import Dict, Any, List, Optional


def read_csv(filepath: str, delimiter: str = ",", limit: int = 1000) -> Dict[str, Any]:
    """
    Read a CSV file
    
    Args:
        filepath: Path to CSV file
        delimiter: Column delimiter (default: comma)
        limit: Maximum rows to read (default: 1000)
    """
    try:
        import pandas as pd
        
        filepath = os.path.expanduser(filepath)
        
        if not os.path.exists(filepath):
            return {"success": False, "error": f"File not found: {filepath}"}
        
        df = pd.read_csv(filepath, delimiter=delimiter, nrows=limit)
        
        return {
            "success": True,
            "filepath": filepath,
            "columns": df.columns.tolist(),
            "row_count": len(df),
            "data": df.to_dict(orient='records'),
            "preview": df.head(10).to_string()
        }
        
    except ImportError:
        # Fallback without pandas
        import csv
        
        filepath = os.path.expanduser(filepath)
        if not os.path.exists(filepath):
            return {"success": False, "error": f"File not found: {filepath}"}
        
        with open(filepath, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f, delimiter=delimiter)
            rows = []
            for i, row in enumerate(reader):
                if i >= limit:
                    break
                rows.append(row)
        
        return {
            "success": True,
            "filepath": filepath,
            "columns": list(rows[0].keys()) if rows else [],
            "row_count": len(rows),
            "data": rows
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def write_csv(filepath: str, data: List[Dict], columns: List[str] = None) -> Dict[str, Any]:
    """
    Write data to a CSV file
    
    Args:
        filepath: Output path
        data: List of dictionaries (rows)
        columns: Optional column order
    """
    try:
        import pandas as pd
        
        filepath = os.path.expanduser(filepath)
        os.makedirs(os.path.dirname(filepath) or ".", exist_ok=True)
        
        df = pd.DataFrame(data)
        if columns:
            df = df[columns]
        
        df.to_csv(filepath, index=False)
        
        return {
            "success": True,
            "filepath": os.path.abspath(filepath),
            "rows_written": len(data)
        }
        
    except ImportError:
        import csv
        
        filepath = os.path.expanduser(filepath)
        os.makedirs(os.path.dirname(filepath) or ".", exist_ok=True)
        
        if not data:
            return {"success": False, "error": "No data to write"}
        
        cols = columns or list(data[0].keys())
        
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=cols)
            writer.writeheader()
            writer.writerows(data)
        
        return {
            "success": True,
            "filepath": os.path.abspath(filepath),
            "rows_written": len(data)
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def read_excel(filepath: str, sheet_name: str = None, limit: int = 1000) -> Dict[str, Any]:
    """
    Read an Excel file
    
    Args:
        filepath: Path to Excel file (.xlsx, .xls)
        sheet_name: Specific sheet to read (default: first sheet)
        limit: Maximum rows to read
    """
    try:
        import pandas as pd
        
        filepath = os.path.expanduser(filepath)
        
        if not os.path.exists(filepath):
            return {"success": False, "error": f"File not found: {filepath}"}
        
        # Get sheet names
        xl = pd.ExcelFile(filepath)
        sheet_names = xl.sheet_names
        
        # Read specified or first sheet
        target_sheet = sheet_name or sheet_names[0]
        df = pd.read_excel(filepath, sheet_name=target_sheet, nrows=limit)
        
        return {
            "success": True,
            "filepath": filepath,
            "sheet_names": sheet_names,
            "current_sheet": target_sheet,
            "columns": df.columns.tolist(),
            "row_count": len(df),
            "data": df.to_dict(orient='records'),
            "preview": df.head(10).to_string()
        }
        
    except ImportError:
        return {"success": False, "error": "pandas and openpyxl required. Run: pip install pandas openpyxl"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def write_excel(filepath: str, data: List[Dict], sheet_name: str = "Sheet1") -> Dict[str, Any]:
    """Write data to an Excel file"""
    try:
        import pandas as pd
        
        filepath = os.path.expanduser(filepath)
        os.makedirs(os.path.dirname(filepath) or ".", exist_ok=True)
        
        df = pd.DataFrame(data)
        df.to_excel(filepath, sheet_name=sheet_name, index=False)
        
        return {
            "success": True,
            "filepath": os.path.abspath(filepath),
            "sheet": sheet_name,
            "rows_written": len(data)
        }
        
    except ImportError:
        return {"success": False, "error": "pandas and openpyxl required"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def read_word(filepath: str) -> Dict[str, Any]:
    """Read a Word document (.docx)"""
    try:
        from docx import Document
        
        filepath = os.path.expanduser(filepath)
        
        if not os.path.exists(filepath):
            return {"success": False, "error": f"File not found: {filepath}"}
        
        doc = Document(filepath)
        
        # Extract text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Extract text from tables
        tables_data = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                table_rows.append([cell.text for cell in row.cells])
            tables_data.append(table_rows)
        
        full_text = "\n\n".join(paragraphs)
        
        return {
            "success": True,
            "filepath": filepath,
            "paragraphs": len(paragraphs),
            "tables": len(tables_data),
            "content": full_text,
            "tables_data": tables_data
        }
        
    except ImportError:
        return {"success": False, "error": "python-docx required. Run: pip install python-docx"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def write_word(filepath: str, content: str, title: str = None) -> Dict[str, Any]:
    """Write content to a Word document"""
    try:
        from docx import Document
        from docx.shared import Pt
        
        filepath = os.path.expanduser(filepath)
        os.makedirs(os.path.dirname(filepath) or ".", exist_ok=True)
        
        doc = Document()
        
        if title:
            doc.add_heading(title, 0)
        
        # Add paragraphs (split by double newlines)
        for para in content.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        
        doc.save(filepath)
        
        return {
            "success": True,
            "filepath": os.path.abspath(filepath)
        }
        
    except ImportError:
        return {"success": False, "error": "python-docx required"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def read_json(filepath: str) -> Dict[str, Any]:
    """Read a JSON file"""
    try:
        filepath = os.path.expanduser(filepath)
        
        if not os.path.exists(filepath):
            return {"success": False, "error": f"File not found: {filepath}"}
        
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        return {
            "success": True,
            "filepath": filepath,
            "data": data,
            "type": type(data).__name__
        }
        
    except json.JSONDecodeError as e:
        return {"success": False, "error": f"Invalid JSON: {e}"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def write_json(filepath: str, data: Any, indent: int = 2) -> Dict[str, Any]:
    """Write data to a JSON file"""
    try:
        filepath = os.path.expanduser(filepath)
        os.makedirs(os.path.dirname(filepath) or ".", exist_ok=True)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=indent, ensure_ascii=False)
        
        return {
            "success": True,
            "filepath": os.path.abspath(filepath)
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}


def convert_csv_to_json(csv_path: str, json_path: str) -> Dict[str, Any]:
    """Convert a CSV file to JSON"""
    try:
        result = read_csv(csv_path)
        if not result["success"]:
            return result
        
        return write_json(json_path, result["data"])
        
    except Exception as e:
        return {"success": False, "error": str(e)}


def convert_json_to_csv(json_path: str, csv_path: str) -> Dict[str, Any]:
    """Convert a JSON file to CSV (must be array of objects)"""
    try:
        result = read_json(json_path)
        if not result["success"]:
            return result
        
        data = result["data"]
        if not isinstance(data, list):
            return {"success": False, "error": "JSON must be an array of objects"}
        
        return write_csv(csv_path, data)
        
    except Exception as e:
        return {"success": False, "error": str(e)}
